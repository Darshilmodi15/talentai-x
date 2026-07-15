"""
Agent 1: Resume Parser
- Detects layout (single-col, two-col, table, image-based)
- Routes to correct text extractor
- Runs 4 specialized LLM prompts in PARALLEL
- Detects AI-written content
- Detects language
"""
import asyncio
import json
import time
import random
from datetime import datetime, timezone
from typing import Optional, Any
import io
import logging

from app.core.pipeline_state import PipelineState, AgentTrace
from app.core.config import settings

import google.generativeai as genai

logger = logging.getLogger(__name__)

logger.warning("PARSE_AGENT_VERSION_2026_06_06_v3_QUOTA_FIX")

# Patterns that indicate Gemini quota/rate-limit errors
QUOTA_ERROR_PATTERNS = [
    "429",
    "too many requests",
    "resource exhausted",
    "resourceexhausted",
    "quota exceeded",
    "rate limit",
    "ratelimit",
    "generaterequestsperday",
]
# ──────────────────────────────────────────────────────────────────
# PROMPTS — tightly scoped per extraction task
# ──────────────────────────────────────────────────────────────────

PROMPT_EXTRACT_ALL = """Extract ALL structured information from this resume text.
Return ONLY valid JSON, no explanation, no markdown.

Schema:
{
  "name": string or null,
  "email": string or null,
  "phone": string or null,
  "location": string or null,
  "summary": string or null,
  "linkedin_url": string or null,
  "github_url": string or null,
  "portfolio_url": string or null,
  "other_urls": [string],
  "experience": [
    {
      "company": string or null,
      "role": string or null,
      "start": string or null,
      "end": string or null,
      "duration_months": integer (estimate if not stated),
      "bullets": [string],
      "skills_mentioned": [string]
    }
  ],
  "experience_months_total": integer,
  "education": [
    {
      "institution": string or null,
      "degree": string or null,
      "field": string or null,
      "year": integer or null,
      "gpa": float or null
    }
  ],
  "skills": [string],
  "certifications": [
    {"name": string, "issuer": string or null, "year": integer or null}
  ],
  "projects": [
    {"name": string, "description": string, "tech": [string], "url": string or null}
  ],
  "publications": [string]
}

If a field is not present, use null. Never invent data.

Resume text:
{text}"""


# ──────────────────────────────────────────────────────────────────
# Layout Detection
# ──────────────────────────────────────────────────────────────────

def detect_pdf_layout(file_bytes: bytes) -> str:
    """Classify PDF layout to route to correct extractor."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                return "image_based"

            page = pdf.pages[0]
            words = page.extract_words()

            if not words:
                return "image_based"

            # Check for two-column layout by x-coordinate bimodal distribution
            x_coords = [w["x0"] for w in words]
            page_mid = page.width / 2

            left_count = sum(1 for x in x_coords if x < page_mid * 0.6)
            right_count = sum(1 for x in x_coords if x > page_mid * 1.4)

            if left_count > 30 and right_count > 30:
                return "two_column"

            # Check for table-heavy
            tables = page.extract_tables()
            if tables and len(tables) > 2:
                return "table_heavy"

            return "single_column"
    except Exception:
        return "single_column"


# ──────────────────────────────────────────────────────────────────
# Text Extractors per Layout
# ──────────────────────────────────────────────────────────────────

def extract_single_column(file_bytes: bytes) -> str:
    import pdfplumber
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        return "\n".join(
            page.extract_text() or "" for page in pdf.pages
        )


def extract_two_column(file_bytes: bytes) -> str:
    """Split page into left/right streams and extract each."""
    texts = []
    import pdfplumber
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            mid = page.width / 2
            left = page.crop((0, 0, mid, page.height)).extract_text() or ""
            right = page.crop((mid, 0, page.width, page.height)).extract_text() or ""
            texts.append(left + "\n" + right)
    return "\n\n".join(texts)


def extract_table_heavy(file_bytes: bytes) -> str:
    """Extract tables first, then remaining text."""
    texts = []
    import pdfplumber
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            table_text = ""
            for table in tables:
                for row in table:
                    table_text += " | ".join((cell or "") for cell in row) + "\n"
            body_text = page.extract_text() or ""
            texts.append(table_text + "\n" + body_text)
    return "\n\n".join(texts)


def extract_with_ocr(file_bytes: bytes) -> str:
    """For scanned/image PDFs — uses PyMuPDF + Tesseract."""
    try:
        import pytesseract
        from PIL import Image
        import fitz  # PyMuPDF fallback

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        texts = []
        for page in doc:
            mat = fitz.Matrix(2, 2)  # 2x zoom for better OCR
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            text = pytesseract.image_to_string(img)
            texts.append(text)
        return "\n\n".join(texts)
    except Exception as e:
        return f"OCR failed: {e}"


def extract_docx(file_bytes: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def extract_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


EXTRACTORS = {
    "single_column": extract_single_column,
    "two_column": extract_two_column,
    "table_heavy": extract_table_heavy,
    "image_based": extract_with_ocr,
}


def get_raw_text(file_bytes: bytes, file_type: str, layout: str) -> str:
    if file_type == "pdf":
        extractor = EXTRACTORS.get(layout, extract_single_column)
        return extractor(file_bytes)
    elif file_type == "docx":
        return extract_docx(file_bytes)
    else:
        return extract_txt(file_bytes)


# ──────────────────────────────────────────────────────────────────
# AI Content Detection
# ──────────────────────────────────────────────────────────────────

def detect_ai_content(text: str) -> float:
    """
    Heuristic AI-content detector.
    AI text has unnaturally uniform sentence lengths (low burstiness).
    Returns probability 0.0–1.0.
    """
    try:
        import re
        sentences = [s.strip() for s in re.split(r'[.!?]+', text) if len(s.strip()) > 10]
        if len(sentences) < 5:
            return 0.0

        lengths = [len(s.split()) for s in sentences]
        mean = sum(lengths) / len(lengths)
        variance = sum((l - mean) ** 2 for l in lengths) / len(lengths)
        std = variance ** 0.5

        burstiness = std / mean if mean > 0 else 1.0
        # Low burstiness = suspicious
        ai_prob = max(0.0, min(1.0, 1.0 - (burstiness / 0.5)))
        return round(ai_prob, 3)
    except Exception:
        return 0.0


# ──────────────────────────────────────────────────────────────────
# LLM Calls — run in parallel
# ──────────────────────────────────────────────────────────────────

class GeminiCallError(Exception):
    def __init__(self, message, raw_response, cleaned_response, traceback_str, exception_type):
        super().__init__(message)
        self.raw_response = raw_response
        self.cleaned_response = cleaned_response
        self.traceback_str = traceback_str
        self.exception_type = exception_type


class GeminiQuotaError(GeminiCallError):
    """Raised specifically for 429 / quota / rate-limit errors from Gemini.
    These should NEVER be silently swallowed."""
    pass


class GeminiNotFoundError(GeminiCallError):
    """Raised for 404 Model Not Found errors."""
    pass


class GeminiPermissionDeniedError(GeminiCallError):
    """Raised for 401/403 provider access problems.

    This must not leak provider tracebacks to the UI. We can still provide a
    deterministic parse fallback from extracted resume text.
    """
    pass


class GeminiJSONError(GeminiCallError):
    """Raised when JSON parsing fails completely even after repair."""
    pass


def _is_quota_error(exception: Exception) -> bool:
    """Check if an exception is a Gemini quota/rate-limit error."""
    error_str = str(exception).lower()
    return any(pattern in error_str for pattern in QUOTA_ERROR_PATTERNS)


def _is_not_found_error(exception: Exception) -> bool:
    error_str = str(exception).lower()
    return any(pattern in error_str for pattern in ["404", "not found", "is no longer available"])


def _is_permission_denied_error(exception: Exception) -> bool:
    error_str = str(exception).lower()
    return any(pattern in error_str for pattern in ["403", "permission denied", "denied access", "unauthorized", "invalid api key"])


async def call_gemini(prompt: str, state: PipelineState, max_tokens: int = 2500, _retries: int = 4) -> dict:
    """Single async Gemini call with exponential backoff for 429 errors.
    Returns parsed JSON or raises GeminiQuotaError / GeminiCallError."""
    import traceback
    content = ""
    cleaned_response = ""
    last_exception = None

    job_id = state.get("job_id", "unknown")
    BACKOFFS = [5, 15, 30]

    for attempt in range(_retries):
        try:
            state["gemini_api_calls"] = (state.get("gemini_api_calls") or 0) + 1
            genai.configure(api_key=settings.GEMINI_API_KEY)
            model = genai.GenerativeModel(settings.GEMINI_MODEL)

            api_key_prefix = settings.GEMINI_API_KEY[:6] + "..." if settings.GEMINI_API_KEY else "MISSING_KEY"
            estimated_tokens = len(prompt) // 4
            logger.info(
                f"Calling Gemini | agent=parse_agent | job_id={job_id} | attempt={attempt+1}/{_retries} | "
                f"key_prefix={api_key_prefix} | model={settings.GEMINI_MODEL} | est_tokens_sent={estimated_tokens}"
            )

            from google.generativeai.types import HarmCategory, HarmBlockThreshold
            safety_settings = {
                HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
            }

            import time
            start_gemini = time.time()
            response = await model.generate_content_async(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    max_output_tokens=8192,
                    temperature=0.1,
                    response_mime_type="application/json",
                ),
                safety_settings=safety_settings,
            )
            latency_ms = int((time.time() - start_gemini) * 1000)
            
            # Extract usage metadata
            try:
                usage = response.usage_metadata
                prompt_tokens = usage.prompt_token_count if usage else 0
                out_tokens = usage.candidates_token_count if usage else 0
            except Exception:
                prompt_tokens = 0
                out_tokens = 0

            logger.info(
                f"Gemini call complete | latency={latency_ms}ms | "
                f"prompt_tokens={prompt_tokens} | output_tokens={out_tokens} | "
            )
            logger.info(f"Gemini candidate count: {len(response.candidates)}")
            if response.candidates:
                candidate = response.candidates[0]
                reason = candidate.finish_reason
                logger.info(f"Gemini finish reason: {reason}")
                
                if reason == 3 or str(reason) == "FinishReason.SAFETY":
                    logger.error("Candidate blocked by SAFETY filters!")
                    if candidate.safety_ratings:
                        for rating in candidate.safety_ratings:
                            logger.error(f"Safety Rating -> Category: {rating.category}, Probability: {rating.probability}, Blocked: {rating.blocked}")
                
            try:
                content = response.text.strip()
                logger.info(f"Gemini response length: {len(content)}")
            except ValueError as ve:
                logger.error(f"Failed to read response.text. Blocked? {ve}")
                content = ""
            
            logger.info(f"Raw Gemini response (first 500): {content[:500]}")

            # Remove markdown wrappers
            cleaned_response = content
            import re
            # Try standard markdown fence first
            fence_match = re.search(r'```(?:json)?\s*(.*?)\s*```', content, re.DOTALL)
            if fence_match:
                cleaned_response = fence_match.group(1).strip()
            else:
                # Try to find JSON object boundaries directly
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    cleaned_response = json_match.group(0).strip()
                else:
                    # Last resort: strip backticks
                    cleaned_response = content.strip("`").strip()
                    if cleaned_response.startswith("json"):
                        cleaned_response = cleaned_response[4:].strip()

            logger.info(f"Cleaned response (first 500): {cleaned_response[:500]}")

            try:
                parsed = json.loads(cleaned_response)
            except Exception as e:
                logger.warning(f"Initial JSON parse failed. Attempting auto-repair. Error: {str(e)}")
                try:
                    import json_repair
                    repaired = json_repair.repair_json(cleaned_response, return_objects=True)
                    if isinstance(repaired, dict):
                        parsed = repaired
                    else:
                        raise ValueError("Repaired output is not a JSON object")
                except Exception as repair_e:
                    logger.exception("JSON auto-repair failed")
                    raise GeminiJSONError(
                        f"INVALID_GEMINI_JSON\n"
                        f"Exception: {type(e).__name__}: {e}\n"
                        f"Repair Exception: {repair_e}\n"
                        f"Raw response length: {len(content)}\n"
                        f"Cleaned response (first 500 chars): {cleaned_response[:500]}",
                        raw_response=content,
                        cleaned_response=cleaned_response,
                        traceback_str=str(repair_e),
                        exception_type=type(repair_e).__name__
                    )

            if not isinstance(parsed, dict):
                raise GeminiJSONError(
                    "Gemini did not return a JSON object",
                    raw_response=content,
                    cleaned_response=cleaned_response,
                    traceback_str="",
                    exception_type="ValueError"
                )

            return parsed

        except GeminiJSONError as ge:
            if attempt < _retries - 1:
                logger.warning(f"Gemini JSON invalid. Retrying ({attempt+1}/{_retries}) with stricter prompt.")
                prompt += "\n\nCRITICAL: You must return ONLY valid JSON. Do not return incomplete or truncated strings."
                continue
            logger.error("Gemini JSON invalid after all retries, bubbling up.")
            raise ge
        except Exception as e:
            last_exception = e
            logger.error(f"CALL_GEMINI_FAILED (attempt {attempt + 1}/{_retries}): {type(e).__name__}: {e}")

            if _is_quota_error(e):
                logger.error("Gemini quota exceeded. Aborting without retry.")
                raise GeminiQuotaError(
                    f"GEMINI_QUOTA_EXHAUSTED\n"
                    f"TYPE={type(e).__name__}\nERROR={str(e)}\n"
                    f"TRACEBACK={traceback.format_exc()}",
                    raw_response=content,
                    cleaned_response=cleaned_response,
                    traceback_str=traceback.format_exc(),
                    exception_type=type(e).__name__,
                )
            
            if _is_not_found_error(e):
                logger.error(f"Gemini model {settings.GEMINI_MODEL} not found (404). Aborting without retry.")
                raise GeminiNotFoundError(
                    f"GEMINI_MODEL_NOT_FOUND\n"
                    f"TYPE={type(e).__name__}\nERROR={str(e)}\n"
                    f"TRACEBACK={traceback.format_exc()}",
                    raw_response=content,
                    cleaned_response=cleaned_response,
                    traceback_str=traceback.format_exc(),
                    exception_type=type(e).__name__,
                )

            if _is_permission_denied_error(e):
                logger.error("Gemini permission denied. Aborting retries and using caller fallback if available.")
                raise GeminiPermissionDeniedError(
                    "AI provider access denied",
                    raw_response=content,
                    cleaned_response=cleaned_response,
                    traceback_str="",
                    exception_type=type(e).__name__,
                )
            
            error_str = str(e).lower()
            is_network = any(p in error_str for p in ["connection", "timeout", "transient", "network", "ssl"])
            if is_network and attempt < len(BACKOFFS):
                backoff = BACKOFFS[attempt]
                logger.warning(f"Network error. Retrying in {backoff}s...")
                await asyncio.sleep(backoff)
                continue
            else:
                raise GeminiCallError(
                    f"CALL_GEMINI_FAILED\nTYPE={type(e).__name__}\nERROR={str(e)}\n"
                    f"TRACEBACK={traceback.format_exc()}",
                    raw_response=content,
                    cleaned_response=cleaned_response,
                    traceback_str=traceback.format_exc(),
                    exception_type=type(e).__name__,
                )

    # Should never reach here, but just in case
    raise GeminiCallError(
        f"CALL_GEMINI_FAILED after {_retries} attempts",
        raw_response=content,
        cleaned_response=cleaned_response,
        traceback_str="",
        exception_type=type(last_exception).__name__ if last_exception else "Unknown",
    )




async def extract_all_info(raw_text: str, state: PipelineState) -> dict:
    """Run a single extraction prompt to conserve quotas."""
    job_id = state.get("job_id", "unknown")
    text_chunk = raw_text[:8000]
    logger.info(f"Starting Gemini extraction, text chunk length={len(text_chunk)}, job_id={job_id}")
    
    return await call_gemini(PROMPT_EXTRACT_ALL.replace("{text}", text_chunk), state, max_tokens=2500)


def heuristic_parse_resume(raw_text: str) -> dict:
    """Best-effort parser used when the AI provider is unavailable/denied.

    It intentionally extracts only evidence present in the resume text. This
    avoids total product failure and lets later normalization/matching still work
    with lower confidence.
    """
    import re
    from app.agents.normalize_agent import SYNONYMS

    text = raw_text or ""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    email_match = re.search(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", text, re.I)
    phone_match = re.search(r"(?:\+?\d[\d\s().-]{7,}\d)", text)
    urls = re.findall(r"https?://[^\s)>,]+|(?:www\.)[^\s)>,]+", text, re.I)

    name = None
    for line in lines[:12]:
        lower = line.lower()
        if any(token in lower for token in ["@", "http", "www.", "resume", "curriculum", "linkedin", "github"]):
            continue
        words = re.findall(r"[A-Za-z][A-Za-z'.-]+", line)
        if 1 <= len(words) <= 4 and len(line) <= 80:
            name = " ".join(words)
            break

    found_skills: list[str] = []
    text_lower = text.lower()
    for alias, canonical in sorted(SYNONYMS.items(), key=lambda kv: len(kv[0]), reverse=True):
        pattern = r"(?<![\w.+#-])" + re.escape(alias.lower()) + r"(?![\w.+#-])"
        if re.search(pattern, text_lower) and canonical not in found_skills:
            found_skills.append(canonical)

    linkedin_url = next((u for u in urls if "linkedin.com" in u.lower()), None)
    github_url = next((u for u in urls if "github.com" in u.lower()), None)

    exp_years = 0
    exp_matches = [int(y) for y in re.findall(r"(\d{1,2})\+?\s*(?:years|yrs)\b", text_lower)]
    if exp_matches:
        exp_years = max(exp_matches)

    return {
        "name": name,
        "email": email_match.group(0) if email_match else None,
        "phone": phone_match.group(0).strip() if phone_match else None,
        "location": None,
        "summary": "Best-effort parse generated without AI provider access.",
        "linkedin_url": linkedin_url,
        "github_url": github_url,
        "portfolio_url": None,
        "other_urls": [u for u in urls if u not in {linkedin_url, github_url}],
        "experience": [],
        "experience_months_total": exp_years * 12,
        "education": [],
        "skills": found_skills,
        "certifications": [],
        "projects": [],
        "publications": [],
    }


def compute_parse_confidence(parsed: dict) -> float:
    """Score how complete the parse is. 0.0–1.0"""
    key_fields = ["name", "email", "skills", "experience", "education"]
    filled = sum(1 for f in key_fields if parsed.get(f))
    return round(filled / len(key_fields), 2)


# ──────────────────────────────────────────────────────────────────
# Main Agent Function
# ──────────────────────────────────────────────────────────────────

async def parse_agent(state: PipelineState) -> PipelineState:
    """
    Agent 1: Parse resume from raw bytes.
    Writes: layout_type, parsed, parse_confidence, resume_language,
            ai_content_probability, traces
    """
    started = time.time()
    retry_count = 0
    error_msg = None
    traceback_str = None
    exception_type = None
    raw_gemini_response = None
    cleaned_gemini_response = None
    raw_text = ""
    status = "success"

    try:
        logger.info(f"=== PARSE_AGENT START === job_id={state.get('job_id')} file={state.get('file_name')}")

        # Step 1: Detect layout for PDFs
        layout = "single_column"
        if state["file_type"] == "pdf":
            layout = detect_pdf_layout(state["raw_file"])
        logger.info(f"Layout detected: {layout}")

        state["layout_type"] = layout

        # Step 2: Extract raw text
        raw_text = get_raw_text(state["raw_file"], state["file_type"], layout)
        logger.info(f"PDF text extraction: {len(raw_text)} chars")
        logger.info(f"Extracted text (first 500 chars): {raw_text[:500]}")

        if not raw_text.strip():
            raise ValueError("No text could be extracted from the file")

        # Step 3: Detect language
        try:
            from langdetect import detect as detect_language
            lang = detect_language(raw_text[:500])
        except Exception:
            lang = "en"
        state["resume_language"] = lang
        logger.info(f"Language detected: {lang}")

        # Step 4: Detect AI content
        state["ai_content_probability"] = detect_ai_content(raw_text)

        # Step 5: Single LLM extraction (with deterministic fallback when no key is configured)
        if not settings.GEMINI_API_KEY or settings.GEMINI_API_KEY == "your_gemini_api_key_here":
            logger.warning("Gemini API key missing/placeholder; using heuristic resume parser")
            parsed = heuristic_parse_resume(raw_text)
            status = "degraded"
            error_msg = "AI provider not configured; used best-effort parser"
        else:
            parsed = await extract_all_info(raw_text, state)

        # Retry if basic info is completely empty (non-quota failures only)
        if not parsed.get("name") and retry_count < 1:
            retry_count += 1
            logger.warning("Basic info empty, retrying extraction...")
            parsed = await extract_all_info(raw_text, state)

        logger.info(
            f"Merged parse result: name={parsed.get('name')}, "
            f"email={parsed.get('email')}, "
            f"phone={parsed.get('phone')}, "
            f"skills_count={len(parsed.get('skills', []))}, "
            f"experience_count={len(parsed.get('experience', []))}, "
            f"education_count={len(parsed.get('education', []))}"
        )

        state["parsed"] = parsed
        confidence = compute_parse_confidence(parsed)
        state["parse_confidence"] = confidence
        state["experience_months_total"] = parsed.get("experience_months_total", 0)

        # Validate: if confidence is 0, the parse is effectively empty
        if confidence == 0.0:
            logger.error(
                "Parse confidence is 0.0 — all key fields are empty. "
                "Marking as failed despite no exception."
            )
            status = "failed"
            error_msg = (
                "Parse completed but extracted zero data. "
                "All key fields (name, email, skills, experience, education) are empty. "
                "This likely indicates a Gemini API issue or incompatible resume format."
            )
            state["errors"] = state.get("errors", []) + [f"parse_agent: {error_msg}"]
        else:
            status = status if status == "degraded" else "success"
            logger.info(f"Parse succeeded with confidence={confidence}, status={status}")

    except GeminiPermissionDeniedError as e:
        logger.error("=== PARSE_AGENT PROVIDER ACCESS DENIED === %s", str(e))
        parsed = heuristic_parse_resume(raw_text)
        confidence = compute_parse_confidence(parsed)
        if confidence > 0.0:
            state["parsed"] = parsed
            state["parse_confidence"] = confidence
            state["experience_months_total"] = parsed.get("experience_months_total", 0)
            state["layout_type"] = state.get("layout_type", "unknown")
            state["resume_language"] = state.get("resume_language", "en")
            state["ai_content_probability"] = state.get("ai_content_probability", 0.0)
            error_msg = "AI provider access denied; used best-effort parser"
            status = "degraded"
        else:
            error_msg = "AI provider access denied. Please check the Gemini API key/project access."
            state["errors"] = state.get("errors", []) + [error_msg]
            state["parsed"] = {}
            state["parse_confidence"] = 0.0
            state["layout_type"] = state.get("layout_type", "unknown")
            state["resume_language"] = state.get("resume_language", "en")
            state["ai_content_probability"] = state.get("ai_content_probability", 0.0)
            state["experience_months_total"] = 0
            status = "failed"
        traceback_str = None
        exception_type = e.exception_type
        raw_gemini_response = None
        cleaned_gemini_response = None

    except GeminiQuotaError as e:
        import traceback as tb
        error_msg = "Gemini quota exceeded"
        logger.error(f"=== PARSE_AGENT QUOTA ERROR === {str(e)}")
        state["errors"] = state.get("errors", []) + [error_msg]
        state["parsed"] = {}
        state["parse_confidence"] = 0.0
        state["layout_type"] = state.get("layout_type", "unknown")
        state["resume_language"] = state.get("resume_language", "en")
        state["ai_content_probability"] = state.get("ai_content_probability", 0.0)
        state["experience_months_total"] = 0
        status = "failed"

        # Do NOT log these to the state traces to avoid leaking to frontend
        traceback_str = None
        exception_type = None
        raw_gemini_response = None
        cleaned_gemini_response = None
        
    except GeminiNotFoundError as e:
        import traceback as tb
        error_msg = "Gemini model not found"
        logger.error(f"=== PARSE_AGENT MODEL ERROR === {str(e)}")
        state["errors"] = state.get("errors", []) + [error_msg]
        state["parsed"] = {}
        state["parse_confidence"] = 0.0
        state["layout_type"] = state.get("layout_type", "unknown")
        state["resume_language"] = state.get("resume_language", "en")
        state["ai_content_probability"] = state.get("ai_content_probability", 0.0)
        state["experience_months_total"] = 0
        status = "failed"

        # Do NOT log these to the state traces to avoid leaking to frontend
        traceback_str = None
        exception_type = None
        raw_gemini_response = None
        cleaned_gemini_response = None
        
    except GeminiJSONError as e:
        import traceback as tb
        error_msg = "Gemini JSON format invalid"
        logger.error(f"=== PARSE_AGENT JSON ERROR === {str(e)}")
        state["errors"] = state.get("errors", []) + [error_msg]
        state["parsed"] = {}
        state["parse_confidence"] = 0.0
        state["layout_type"] = state.get("layout_type", "unknown")
        state["resume_language"] = state.get("resume_language", "en")
        state["ai_content_probability"] = state.get("ai_content_probability", 0.0)
        state["experience_months_total"] = 0
        status = "failed"

        traceback_str = None
        exception_type = e.exception_type
        raw_gemini_response = None
        cleaned_gemini_response = None

    except GeminiCallError as e:
        logger.error("=== PARSE_AGENT AI PROVIDER ERROR === %s", str(e))
        parsed = heuristic_parse_resume(raw_text)
        confidence = compute_parse_confidence(parsed)
        if confidence > 0.0:
            state["parsed"] = parsed
            state["parse_confidence"] = confidence
            state["experience_months_total"] = parsed.get("experience_months_total", 0)
            state["layout_type"] = state.get("layout_type", "unknown")
            state["resume_language"] = state.get("resume_language", "en")
            state["ai_content_probability"] = state.get("ai_content_probability", 0.0)
            error_msg = "AI provider unavailable; used best-effort parser"
            status = "degraded"
        else:
            error_msg = "AI provider unavailable. Please try again later or check provider configuration."
            state["errors"] = state.get("errors", []) + [error_msg]
            state["parsed"] = {}
            state["parse_confidence"] = 0.0
            state["layout_type"] = state.get("layout_type", "unknown")
            state["resume_language"] = state.get("resume_language", "en")
            state["ai_content_probability"] = state.get("ai_content_probability", 0.0)
            state["experience_months_total"] = 0
            status = "failed"
        traceback_str = None
        exception_type = e.exception_type
        raw_gemini_response = None
        cleaned_gemini_response = None

    except Exception as e:
        import traceback as tb
        error_msg = str(e)
        logger.error(f"=== PARSE_AGENT ERROR === {type(e).__name__}: {error_msg}")
        state["errors"] = state.get("errors", []) + [f"parse_agent: {error_msg}"]
        state["parsed"] = {}
        state["parse_confidence"] = 0.0
        state["layout_type"] = state.get("layout_type", "unknown")
        state["resume_language"] = state.get("resume_language", "en")
        state["ai_content_probability"] = state.get("ai_content_probability", 0.0)
        state["experience_months_total"] = 0
        status = "failed"

        traceback_str = getattr(e, "traceback_str", tb.format_exc())
        exception_type = getattr(e, "exception_type", type(e).__name__)
        raw_gemini_response = getattr(e, "raw_response", None)
        cleaned_gemini_response = getattr(e, "cleaned_response", None)

    # Record trace
    fields_extracted = len([v for v in (state.get("parsed") or {}).values() if v])
    trace: AgentTrace = {
        "agent": "parse_agent",
        "started_at": datetime.now(timezone.utc).isoformat(),
        "duration_ms": int((time.time() - started) * 1000),
        "status": status,
        "quality_score": state.get("parse_confidence", 0.0),
        "fields_extracted": fields_extracted,
        "retry_count": retry_count,
        "error": error_msg,
    }

    if status == "failed":
        trace["exception_type"] = exception_type
        trace["traceback"] = traceback_str
        trace["raw_gemini_response"] = raw_gemini_response
        trace["cleaned_gemini_response"] = cleaned_gemini_response

    logger.info(
        f"=== PARSE_AGENT END === status={status} fields_extracted={fields_extracted} "
        f"confidence={state.get('parse_confidence', 0.0)} duration_ms={trace['duration_ms']}"
    )

    state["traces"] = state.get("traces", []) + [trace]
    return state
